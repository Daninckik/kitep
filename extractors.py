# -*- coding: utf-8 -*-
"""Извлечение текста из файлов книг.

extract() возвращает (pages, notes):
  pages — [{"label": "стр. 1", "text": "..."}]
  notes — служебные примечания (например, «текст распознан с фото страниц»)
Для форматов без страниц текст режется на псевдостраницы ~1800 символов.
Если вместо текста в книге фотографии страниц — текст распознаётся через OCR
(Tesseract, языки: русский, кыргызский, английский).
"""
import base64
import io
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from html.parser import HTMLParser

CHUNK = 1800

# ---------- OCR ----------
OCR_MAX_PAGES = 200          # максимум распознаваемых страниц/картинок за один файл
OCR_MIN_CHARS_PER_PAGE = 40  # меньше — считаем, что текстового слоя на странице нет
OCR_MAX_SIDE = 2400          # крупнее не нужно — только замедляет распознавание
OCR_WORKERS = min(8, os.cpu_count() or 2)
_ocr_langs = "rus+kir+eng"

NOTE_OCR = (
    "Вместо части страниц в файле фотографии — текст распознан автоматически (OCR). "
    "Возможны ошибки распознавания, спорные места проверьте по оригиналу."
)


def _prep_image(img):
    """Серый + разумный размер: качество то же, распознавание быстрее."""
    if img.mode != "L":
        img = img.convert("L")
    w, h = img.size
    if max(w, h) > OCR_MAX_SIDE:
        k = OCR_MAX_SIDE / max(w, h)
        img = img.resize((int(w * k), int(h * k)))
    return img


def _ocr_image(img) -> str:
    """PIL.Image -> распознанный текст."""
    global _ocr_langs
    import pytesseract

    img = _prep_image(img)
    config = "-c tessedit_do_invert=0"
    try:
        return pytesseract.image_to_string(img, lang=_ocr_langs, config=config)
    except pytesseract.TesseractError:
        if _ocr_langs != "rus+eng":
            _ocr_langs = "rus+eng"
            return pytesseract.image_to_string(img, lang=_ocr_langs, config=config)
        raise


def _ocr_image_bytes(raw: bytes) -> str:
    from PIL import Image

    with Image.open(io.BytesIO(raw)) as img:
        img.load()
        return _ocr_image(img)


def _ocr_batch(sources: list) -> list:
    """Распознаёт пачку картинок параллельно на всех ядрах.

    sources: список PIL.Image или bytes. Возвращает список текстов той же длины.
    Tesseract — отдельный процесс на картинку, поэтому потоки дают
    полноценное ускорение (OMP_THREAD_LIMIT=1 в образе).
    """
    from concurrent.futures import ThreadPoolExecutor

    def work(src):
        try:
            if isinstance(src, bytes):
                return _ocr_image_bytes(src)
            return _ocr_image(src)
        except Exception:
            return ""

    if len(sources) == 1:
        return [work(sources[0])]
    with ThreadPoolExecutor(max_workers=OCR_WORKERS) as ex:
        return list(ex.map(work, sources))


class _HTMLText(HTMLParser):
    SKIP = {"script", "style"}

    def __init__(self):
        super().__init__()
        self.parts = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in self.SKIP:
            self._skip += 1
        if tag in ("p", "br", "div", "h1", "h2", "h3", "li", "tr"):
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in self.SKIP and self._skip:
            self._skip -= 1

    def handle_data(self, data):
        if not self._skip:
            self.parts.append(data)

    def text(self):
        return re.sub(r"[ \t]+", " ", "".join(self.parts))


def _html_to_text(html: str) -> str:
    p = _HTMLText()
    try:
        p.feed(html)
    except Exception:
        return re.sub(r"<[^>]+>", " ", html)
    return p.text()


def _decode(raw: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1251"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return raw.decode("utf-8", "replace")


def _chunk(text: str, label: str = "фрагмент") -> list:
    text = text.strip()
    if not text:
        return []
    pages = []
    pos = 0
    n = 1
    while pos < len(text):
        end = min(pos + CHUNK, len(text))
        if end < len(text):
            brk = text.rfind(" ", pos + CHUNK - 200, end)
            if brk > pos:
                end = brk
        pages.append({"label": f"{label} {n}", "text": text[pos:end]})
        pos = end
        n += 1
    return pages


def _extract_pdf(path: str, notes: list) -> list:
    from pypdf import PdfReader

    reader = PdfReader(path)
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")

    # страницы-фотографии (нет текстового слоя) распознаём через OCR
    need_ocr = [i for i, t in enumerate(texts) if len(t.strip()) < OCR_MIN_CHARS_PER_PAGE]
    if need_ocr:
        import pypdfium2 as pdfium

        if len(need_ocr) > OCR_MAX_PAGES:
            notes.append(
                f"Фото-страниц больше лимита: распознаны первые {OCR_MAX_PAGES} "
                f"из {len(need_ocr)}."
            )
            need_ocr = need_ocr[:OCR_MAX_PAGES]

        ocr_done = 0
        pdf = pdfium.PdfDocument(path)
        try:
            # рендерим партиями (pdfium не потокобезопасен), OCR — параллельно
            batch_size = OCR_WORKERS * 2
            for b in range(0, len(need_ocr), batch_size):
                batch = need_ocr[b:b + batch_size]
                images = [
                    _prep_image(pdf[i].render(scale=2.0, grayscale=True).to_pil())
                    for i in batch
                ]
                for i, txt in zip(batch, _ocr_batch(images)):
                    if txt.strip():
                        texts[i] = txt
                        ocr_done += 1
        finally:
            pdf.close()
        if ocr_done:
            notes.append(NOTE_OCR)

    return [
        {"label": f"стр. {i}", "text": t}
        for i, t in enumerate(texts, 1)
        if t.strip()
    ]


def _extract_docx(path: str, notes: list) -> list:
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)

    # документ-скан: текста нет, страницы вставлены картинками
    if len(text.strip()) < 200:
        with zipfile.ZipFile(path) as z:
            image_names = sorted(
                n for n in z.namelist()
                if n.startswith("word/media/") and re.search(r"\.(jpe?g|png|tiff?|webp)$", n, re.I)
            )
            raws = [z.read(n) for n in image_names[:OCR_MAX_PAGES]]
        ocr_pages = [
            {"label": f"фото {i}", "text": txt}
            for i, txt in enumerate(_ocr_batch(raws), 1)
            if txt.strip()
        ]
        if ocr_pages:
            notes.append(NOTE_OCR)
            return ocr_pages
    return _chunk(text)


def _extract_fb2(path: str, notes: list) -> list:
    raw = open(path, "rb").read()
    try:
        root = ET.fromstring(raw)
    except ET.ParseError:
        return _chunk(_html_to_text(_decode(raw)))
    ns = ""
    if root.tag.startswith("{"):
        ns = root.tag.split("}")[0] + "}"
    parts = []
    for body in root.iter(f"{ns}body"):
        parts.append(" ".join(body.itertext()))
    if not parts:
        parts = [" ".join(root.itertext())]
    text = "\n".join(parts)

    # книга-скан: текста почти нет, зато есть вложенные картинки страниц
    if len(text.strip()) < 200:
        images = list(root.iter(f"{ns}binary"))
        raws = []
        for b in images[:OCR_MAX_PAGES]:
            try:
                raws.append(base64.b64decode(b.text or ""))
            except Exception:
                continue
        ocr_pages = [
            {"label": f"фото {i}", "text": txt}
            for i, txt in enumerate(_ocr_batch(raws), 1)
            if txt.strip()
        ]
        if ocr_pages:
            notes.append(NOTE_OCR)
            if len(images) > OCR_MAX_PAGES:
                notes.append(
                    f"Фото-страниц больше лимита: распознаны первые {OCR_MAX_PAGES} "
                    f"из {len(images)}."
                )
            return ocr_pages
    return _chunk(text)


def _extract_epub(path: str, notes: list) -> list:
    pages = []
    image_names = []
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist() if re.search(r"\.x?html?$", n, re.I)]
        names.sort()
        for n in names:
            html = _decode(z.read(n))
            txt = _html_to_text(html).strip()
            if txt:
                pages.extend(_chunk(txt, label=f"{n.rsplit('/', 1)[-1]} · фрагмент"))
        # книга-скан: html пустые, страницы лежат картинками
        if sum(len(p["text"]) for p in pages) < 200:
            image_names = sorted(
                n for n in z.namelist() if re.search(r"\.(jpe?g|png|tiff?|webp)$", n, re.I)
            )
            raws = [z.read(n) for n in image_names[:OCR_MAX_PAGES]]
            ocr_pages = [
                {"label": f"фото {i}", "text": txt}
                for i, txt in enumerate(_ocr_batch(raws), 1)
                if txt.strip()
            ]
            if ocr_pages:
                notes.append(NOTE_OCR)
                if len(image_names) > OCR_MAX_PAGES:
                    notes.append(
                        f"Фото-страниц больше лимита: распознаны первые {OCR_MAX_PAGES} "
                        f"из {len(image_names)}."
                    )
                return ocr_pages
    return pages


def _extract_image(path: str, notes: list) -> list:
    txt = _ocr_image_bytes(open(path, "rb").read())
    if txt.strip():
        notes.append(NOTE_OCR)
    return _chunk(txt, label="фото · фрагмент")


def _extract_txt(path: str, notes: list) -> list:
    return _chunk(_decode(open(path, "rb").read()))


def _extract_html_file(path: str, notes: list) -> list:
    return _chunk(_html_to_text(_decode(open(path, "rb").read())))


def _rtf_to_text(data: bytes) -> str:
    """Компактный конвертер RTF -> текст (по алгоритму striprtf)."""
    text = data.decode("latin-1", "ignore")
    pattern = re.compile(
        r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\'([0-9a-f]{2})|\\([^a-z])|([{}])|[\r\n]+|(.)",
        re.I | re.S,
    )
    destinations = {
        "fonttbl", "colortbl", "stylesheet", "info", "pict", "themedata",
        "generator", "operator", "header", "footer", "xmlnstbl", "listtable",
        "listoverridetable", "rsidtbl", "latentstyles", "datastore",
    }
    specials = {"par": "\n", "line": "\n", "tab": "\t", "sect": "\n",
                "page": "\n", "emdash": "—", "endash": "–", "lquote": "'",
                "rquote": "'", "ldblquote": "«", "rdblquote": "»", "~": " "}
    out = []
    stack = []
    ignorable = False
    ucskip = 1
    curskip = 0
    for m in pattern.finditer(text):
        word, arg, hexcode, char, brace, tchar = m.groups()
        if brace:
            if brace == "{":
                stack.append((ucskip, ignorable))
            elif stack:
                ucskip, ignorable = stack.pop()
            curskip = 0
        elif char:
            curskip = 0
            if char == "*":
                ignorable = True
            elif char in "\\{}" and not ignorable:
                out.append(char)
            elif char == "~" and not ignorable:
                out.append(" ")
        elif word:
            curskip = 0
            w = word.lower()
            if w in destinations:
                ignorable = True
            elif ignorable:
                pass
            elif w in specials:
                out.append(specials[w])
            elif w == "uc":
                ucskip = int(arg or 1)
            elif w == "u":
                c = int(arg or 0)
                if c < 0:
                    c += 0x10000
                out.append(chr(c))
                curskip = ucskip
        elif hexcode:
            if curskip > 0:
                curskip -= 1
            elif not ignorable:
                out.append(bytes([int(hexcode, 16)]).decode("cp1251", "ignore"))
        elif tchar:
            if curskip > 0:
                curskip -= 1
            elif not ignorable:
                out.append(tchar)
    return "".join(out)


def _extract_rtf(path: str, notes: list) -> list:
    return _chunk(_rtf_to_text(open(path, "rb").read()))


def _extract_djvu(path: str, notes: list) -> list:
    """DJVU: конвертируем в PDF утилитой ddjvu (djvulibre) и распознаём как PDF."""
    import shutil
    import subprocess

    if not shutil.which("ddjvu"):
        raise ValueError(
            "для DJVU нужна утилита ddjvu (djvulibre). Запустите систему в Docker "
            "или конвертируйте файл в PDF и загрузите снова"
        )
    tmp_pdf = f"{path}.conv.pdf"
    try:
        subprocess.run(
            ["ddjvu", "-format=pdf", "-quality=85", path, tmp_pdf],
            check=True, capture_output=True, timeout=600,
        )
        notes.append("DJVU сконвертирован в PDF автоматически.")
        return _extract_pdf(tmp_pdf, notes)
    except subprocess.CalledProcessError as e:
        raise ValueError(f"не удалось прочитать DJVU: {e.stderr.decode('utf-8', 'ignore')[:200]}")
    finally:
        if os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)


ZIP_INNER_PRIORITY = [
    ".fb2", ".docx", ".rtf", ".html", ".htm", ".txt", ".pdf",
    ".jpg", ".jpeg", ".png", ".tif", ".tiff", ".webp",
]


def _extract_zip(path: str, notes: list) -> list:
    """ZIP-контейнер (например, book.fb2.zip) — достаём книгу изнутри.

    Архив с фотографиями страниц распознаём через OCR постранично.
    """
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist() if not n.endswith("/")]
        image_names = sorted(
            n for n in names if re.search(r"\.(jpe?g|png|tiff?|webp)$", n, re.I)
        )
        doc_names = [n for n in names if n not in set(image_names)]
        for ext in ZIP_INNER_PRIORITY:
            for n in doc_names:
                if n.lower().endswith(ext):
                    tmp = f"{path}.inner{ext}"
                    with open(tmp, "wb") as f:
                        f.write(z.read(n))
                    try:
                        return EXTRACTORS[ext](tmp, notes)
                    finally:
                        os.unlink(tmp)
        if image_names:
            raws = [z.read(n) for n in image_names[:OCR_MAX_PAGES]]
            ocr_pages = [
                {"label": f"фото {i}", "text": txt}
                for i, txt in enumerate(_ocr_batch(raws), 1)
                if txt.strip()
            ]
            if ocr_pages:
                notes.append(NOTE_OCR)
                if len(image_names) > OCR_MAX_PAGES:
                    notes.append(
                        f"Фото-страниц больше лимита: распознаны первые {OCR_MAX_PAGES} "
                        f"из {len(image_names)}."
                    )
                return ocr_pages
    raise ValueError("в ZIP-архиве не нашлось книги поддерживаемого формата")


EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".fb2": _extract_fb2,
    ".epub": _extract_epub,
    ".txt": _extract_txt,
    ".html": _extract_html_file,
    ".htm": _extract_html_file,
    ".rtf": _extract_rtf,
    ".djvu": _extract_djvu,
    ".zip": _extract_zip,
    ".jpg": _extract_image,
    ".jpeg": _extract_image,
    ".png": _extract_image,
    ".tif": _extract_image,
    ".tiff": _extract_image,
    ".webp": _extract_image,
}

SUPPORTED = set(EXTRACTORS)


def extract(path: str, ext: str) -> tuple:
    """-> (pages, notes)"""
    ext = ext.lower()
    if ext not in EXTRACTORS:
        raise ValueError(f"Формат {ext} не поддерживается")
    notes: list = []
    pages = EXTRACTORS[ext](path, notes)
    return pages, notes
